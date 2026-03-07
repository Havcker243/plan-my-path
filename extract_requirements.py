#!/usr/bin/env python3
"""
Extract degree requirements from the Balance Sheet Word document
"""

try:
    from docx import Document
    from docx.table import Table
    import json
    import sys

    # Read the Word document
    doc_path = 'data/Balance Sheet B.S. Computer Science - Oludolapo Adegbesan (1).docx'
    doc = Document(doc_path)

    print("=" * 80)
    print("BALANCE SHEET CONTENT")
    print("=" * 80)
    print()

    # Extract paragraphs
    print("PARAGRAPHS:")
    print("-" * 80)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text}")
    print()

    # Extract tables
    print("\nTABLES:")
    print("-" * 80)
    for i, table in enumerate(doc.tables):
        print(f"\n\nTable {i+1}:")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))

except ImportError:
    print("ERROR: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
