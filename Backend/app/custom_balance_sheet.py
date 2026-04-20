from __future__ import annotations

import io
import re
from typing import Iterable


RE_COURSE_CODE = re.compile(r"\b([A-Z]{2,6})[-\s]?(\d{2,4}[A-Z0-9]*)\b")


def normalize_code(code: str) -> str:
    return re.sub(r"[-\s]+", " ", code.strip().upper())


def scan_balance_sheet_pdf(file_bytes: bytes, course_codes: Iterable[str]) -> dict:
    """
    Locate course-code text in a searchable balance-sheet PDF.

    First tries the PDF text layer. If that finds no rows, falls back to OCR
    when PyMuPDF + pytesseract + the Tesseract binary are available.
    """
    from pypdf import PdfReader

    wanted = {normalize_code(code) for code in course_codes if code.strip()}
    if not wanted:
        return {"method": "text", "confidence": 0, "matches": [], "unmatched_codes": []}

    reader = PdfReader(io.BytesIO(file_bytes))
    matches: list[dict] = []
    found: set[str] = set()

    for page_index, page in enumerate(reader.pages):
        fragments: list[dict] = []

        def visitor_text(text, _cm, tm, _font_dict, font_size):
            value = (text or "").strip()
            if not value:
                return
            fragments.append({
                "text": value,
                "x": float(tm[4]),
                "y": float(tm[5]),
                "font_size": float(font_size or 9),
            })

        try:
            page.extract_text(visitor_text=visitor_text)
        except TypeError:
            continue

        page_matches = _match_text_fragments(page_index, fragments, wanted)
        matches.extend(page_matches)
        found.update(match["course_code"] for match in page_matches)

    if matches:
        return {
            "method": "text",
            "confidence": 1,
            "matches": matches,
            "unmatched_codes": sorted(wanted - found),
        }

    ocr_result = _scan_with_ocr(file_bytes, wanted)
    if ocr_result is not None:
        return ocr_result

    return {
        "method": "text",
        "confidence": 0,
        "matches": [],
        "unmatched_codes": sorted(wanted),
        "warning": "No searchable course rows found. OCR is not available in this environment.",
    }


def _match_text_fragments(page_index: int, fragments: list[dict], wanted: set[str]) -> list[dict]:
    matches: list[dict] = []
    line_buckets: dict[int, list[dict]] = {}
    for fragment in fragments:
        bucket = round(fragment["y"] / 4)
        line_buckets.setdefault(bucket, []).append(fragment)

    for line in line_buckets.values():
        ordered = sorted(line, key=lambda item: item["x"])
        text = " ".join(item["text"] for item in ordered)
        if not text:
            continue
        line_x = min(item["x"] for item in ordered)
        line_y = sum(item["y"] for item in ordered) / len(ordered)
        avg_font = sum(item["font_size"] for item in ordered) / len(ordered)
        avg_char_width = max(4.0, avg_font * 0.52)

        for match in RE_COURSE_CODE.finditer(text):
            code = normalize_code(f"{match.group(1)} {match.group(2)}")
            if code not in wanted:
                continue
            matches.append({
                "course_code": code,
                "page": page_index,
                "x": line_x + match.start() * avg_char_width,
                "y": line_y,
                "font_size": max(8, min(11, avg_font)),
                "line_text": text[:240],
                "confidence": 1,
            })

    return matches


def _scan_with_ocr(file_bytes: bytes, wanted: set[str]) -> dict | None:
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError:
        return None

    matches: list[dict] = []
    found: set[str] = set()
    confidences: list[float] = []

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception:
        return None

    for page_index in range(len(doc)):
        page = doc[page_index]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        except Exception:
            return None

        scale_x = float(page.rect.width) / float(pix.width)
        scale_y = float(page.rect.height) / float(pix.height)
        count = len(data.get("text", []))

        for index in range(count):
            word = str(data["text"][index] or "").strip()
            if not word:
                continue

            candidate = _ocr_course_candidate(data, index, count)
            if not candidate or candidate["code"] not in wanted:
                continue

            left = candidate["left"]
            top = candidate["top"]
            height = candidate["height"]
            confidence = candidate["confidence"]
            x = left * scale_x
            y = (pix.height - top - height) * scale_y

            matches.append({
                "course_code": candidate["code"],
                "page": page_index,
                "x": x,
                "y": y,
                "font_size": 9,
                "line_text": candidate["text"][:240],
                "confidence": confidence,
            })
            found.add(candidate["code"])
            confidences.append(confidence)

    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
    return {
        "method": "ocr",
        "confidence": round(avg_confidence, 3),
        "matches": matches,
        "unmatched_codes": sorted(wanted - found),
    }


def scan_balance_sheet_docx(file_bytes: bytes, course_codes: Iterable[str]) -> dict:
    """
    Extract course codes from a Word (.docx) balance sheet.
    Returns matches in the same shape as scan_balance_sheet_pdf but with
    placeholder coordinates (x=0, y=0) since we generate a new PDF on export.
    """
    try:
        from docx import Document
    except ImportError:
        return {
            "method": "docx",
            "confidence": 0,
            "matches": [],
            "unmatched_codes": sorted(course_codes),
            "warning": "python-docx is not installed on this server.",
        }

    wanted = {normalize_code(code) for code in course_codes if code.strip()}
    if not wanted:
        return {"method": "docx", "confidence": 0, "matches": [], "unmatched_codes": []}

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return {
            "method": "docx",
            "confidence": 0,
            "matches": [],
            "unmatched_codes": sorted(wanted),
            "warning": "Could not open the Word document.",
        }

    found: set[str] = set()
    matches: list[dict] = []
    preview_lines: list[str] = []

    def _add_preview_line(text: str) -> None:
        line = " ".join((text or "").split())
        if line:
            preview_lines.append(line)

    def _scan_text(text: str) -> None:
        _add_preview_line(text)
        for match in RE_COURSE_CODE.finditer(text):
            code = normalize_code(f"{match.group(1)} {match.group(2)}")
            if code in wanted and code not in found:
                found.add(code)
                matches.append({
                    "course_code": code,
                    "page": 0,
                    "x": 0,
                    "y": 0,
                    "font_size": 9,
                    "line_text": text[:240],
                    "confidence": 1,
                })

    for para in doc.paragraphs:
        _scan_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                _scan_text(" | ".join(cells))

    return {
        "method": "docx",
        "confidence": 1 if matches else 0,
        "matches": matches,
        "unmatched_codes": sorted(wanted - found),
        "preview_lines": preview_lines[:250],
    }


def fill_balance_sheet_docx(file_bytes: bytes, rows: Iterable[dict]) -> bytes:
    """
    Add plan markings to the uploaded Word balance sheet and return the edited
    document bytes. This preserves the user's source document instead of
    generating a FiskGrad replacement layout.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    row_map = {
        normalize_code(str(row.get("code", ""))): row
        for row in rows
        if str(row.get("code", "")).strip()
    }
    applied: set[str] = set()

    def _matched_row(text: str) -> tuple[str, dict] | None:
        for match in RE_COURSE_CODE.finditer(text or ""):
            code = normalize_code(f"{match.group(1)} {match.group(2)}")
            row = row_map.get(code)
            if row and code not in applied:
                return code, row
        return None

    def _fill_text_container(paragraphs) -> None:
        for paragraph in paragraphs:
            text = paragraph.text or ""
            match = _matched_row(text)
            if not match:
                continue
            code, row = match
            marker = _docx_marker(row)
            if marker:
                run = paragraph.add_run(f"  {marker}")
                run.bold = True
            applied.add(code)

    _fill_text_container(doc.paragraphs)

    seen_cells: set[int] = set()
    for table in doc.tables:
        header_map = _docx_table_header_map(table)
        for table_row in table.rows:
            row_text = " | ".join(" ".join(cell.text.split()) for cell in table_row.cells if cell.text.strip())
            match = _matched_row(row_text)
            if match:
                code, row = match
                if _fill_docx_table_row(table_row, row, header_map):
                    applied.add(code)
                    continue

            for cell in table_row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                _fill_text_container(cell.paragraphs)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _docx_table_header_map(table) -> dict[str, int]:
    header_map: dict[str, int] = {}
    for table_row in table.rows[:6]:
        for index, cell in enumerate(table_row.cells):
            text = " ".join(cell.text.lower().split())
            if not text:
                continue
            if "grade" in text and "grade" not in header_map:
                header_map["grade"] = index
            if (
                ("semester earned" in text or "term" in text or text == "semester")
                and "term" not in header_map
            ):
                header_map["term"] = index
            if ("comment" in text or "note" in text or "status" in text) and "status" not in header_map:
                header_map["status"] = index
            if (
                ("credit hour" in text or text in {"credits", "credit", "hours"})
                and "credits" not in header_map
            ):
                header_map["credits"] = index
    return header_map


def _fill_docx_table_row(table_row, row: dict, header_map: dict[str, int]) -> bool:
    cells = table_row.cells
    if not cells:
        return False

    code_cell_index = 0
    for index, cell in enumerate(cells):
        if RE_COURSE_CODE.search(cell.text or ""):
            code_cell_index = index
            break

    wrote = False

    def write_mapped(field: str, value: str) -> bool:
        if not value:
            return False
        index = header_map.get(field)
        if index is None or index >= len(cells) or index == code_cell_index:
            return False
        return _write_docx_cell(cells[index], value, append=field == "status")

    grade = str(row.get("grade") or "").strip()
    term = str(row.get("term") or "").strip()
    credits = str(row.get("credits") or "").strip()
    status = _docx_status_label(row)

    wrote = write_mapped("grade", grade) or wrote
    wrote = write_mapped("term", term) or wrote
    wrote = write_mapped("status", status) or wrote

    # Only fill credits when the mapped cell is blank. Most balance sheets
    # already contain official credit-hour values that should stay untouched.
    credits_index = header_map.get("credits")
    if credits and credits_index is not None and credits_index < len(cells):
        credit_cell = cells[credits_index]
        if not credit_cell.text.strip():
            wrote = _write_docx_cell(credit_cell, credits) or wrote

    if wrote:
        return True

    fallback_values = [value for value in [grade, term, status] if value]
    for cell in cells[code_cell_index + 1:]:
        if not fallback_values:
            break
        if cell.text.strip():
            continue
        wrote = _write_docx_cell(cell, fallback_values.pop(0)) or wrote

    if wrote:
        return True

    marker = _docx_marker(row)
    if marker:
        return _append_docx_cell(cells[code_cell_index], marker)
    return False


def _write_docx_cell(cell, value: str, append: bool = False) -> bool:
    value = value.strip()
    if not value:
        return False
    if cell.text.strip():
        if not append:
            return False
        return _append_docx_cell(cell, value)
    cell.text = value
    return True


def _append_docx_cell(cell, value: str) -> bool:
    value = value.strip()
    if not value:
        return False
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(f" {value}")
    run.bold = True
    return True


def _docx_status_label(row: dict) -> str:
    status = str(row.get("status") or "").strip().lower()
    if status == "completed":
        return "Completed"
    if status == "planned":
        return "Planned"
    return ""


def _docx_marker(row: dict) -> str:
    status = str(row.get("status") or "").strip().lower()
    mark = "X" if status == "completed" else "N" if status == "planned" else ""
    parts = [part for part in [
        mark,
        str(row.get("grade") or "").strip(),
        str(row.get("term") or "").strip(),
        str(row.get("credits") or "").strip(),
    ] if part]
    return " | ".join(parts)


def _ocr_course_candidate(data: dict, index: int, count: int) -> dict | None:
    word = str(data["text"][index] or "").strip().upper()
    next_word = str(data["text"][index + 1] or "").strip().upper() if index + 1 < count else ""
    compact = f"{word} {next_word}".strip()
    match = RE_COURSE_CODE.search(compact)
    if not match:
        return None

    code = normalize_code(f"{match.group(1)} {match.group(2)}")
    left = int(data["left"][index])
    top = int(data["top"][index])
    width = int(data["width"][index])
    height = int(data["height"][index])
    confidence_values = []

    for offset in (0, 1):
        idx = index + offset
        if idx >= count:
            continue
        try:
            confidence_values.append(max(0.0, float(data["conf"][idx]) / 100.0))
        except (TypeError, ValueError):
            pass
        if offset == 1 and next_word:
            width = max(width, int(data["left"][idx]) + int(data["width"][idx]) - left)
            top = min(top, int(data["top"][idx]))
            height = max(height, int(data["height"][idx]))

    return {
        "code": code,
        "left": left,
        "top": top,
        "width": width,
        "height": height,
        "confidence": sum(confidence_values) / len(confidence_values) if confidence_values else 0,
        "text": compact,
    }
